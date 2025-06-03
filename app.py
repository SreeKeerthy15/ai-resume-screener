import streamlit as st
from pdfminer.high_level import extract_text
import docx
import spacy
from docx import Document
from io import BytesIO
from fpdf import FPDF
import nltk
from nltk.corpus import wordnet
import openai

# Set your OpenAI API key here or via environment variable
openai.api_key = "YOUR_OPENAI_API_KEY"

nltk.download("wordnet")

nlp = spacy.load("en_core_web_sm")

SKILL_CATEGORIES = {
    "Technical": ["python", "java", "c++", "sql", "html", "css"],
    "Tools": ["git", "docker", "kubernetes", "aws"],
    "Soft Skills": ["communication", "leadership", "teamwork"],
    "Certifications": ["aws certified", "pmp", "azure fundamentals"],
}

def categorize_skill_with_gpt(skill_name: str) -> str:
    prompt = f"""You are a helpful assistant that categorizes a given skill into one of these categories:
(a) Technical
(b) Tools
(c) Soft Skills
(d) Certifications

Skill: '{skill_name}'

Respond ONLY with the letter (a, b, c, or d) corresponding to the most appropriate category."""

    try:
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=1,
            temperature=0,
            n=1,
            stop=None,
        )
        answer = response.choices[0].text.strip().lower()

        mapping = {
            "a": "Technical",
            "b": "Tools",
            "c": "Soft Skills",
            "d": "Certifications"
        }

        return mapping.get(answer, "Technical")

    except Exception as e:
        print(f"OpenAI API error: {e}")
        return "Technical"

def categorize_skills(keywords):
    categorized = {cat: [] for cat in SKILL_CATEGORIES}

    for word in keywords:
        found = False
        for category, skill_list in SKILL_CATEGORIES.items():
            if word in skill_list:
                categorized[category].append(word)
                found = True
                break
        if not found:
            gpt_category = categorize_skill_with_gpt(word)
            categorized[gpt_category].append(word)

    return categorized

def extract_text_from_pdf(file):
    return extract_text(file)

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_keywords(text):
    doc = nlp(text.lower())
    return set([token.text for token in doc if token.is_alpha and not token.is_stop])

def score_meter(score):
    if score < 30:
        color = "red"
    elif score < 70:
        color = "orange"
    else:
        color = "green"

    bar = f"""
    <div style="background-color: lightgray; border-radius: 8px; height: 25px;">
        <div style="width: {score}%; background-color: {color}; height: 100%; border-radius: 8px; text-align: center; color: white;">
            <b>{score:.2f}%</b>
        </div>
    </div>
    """
    st.markdown(bar, unsafe_allow_html=True)

def generate_report(match_score, matched, missing):
    doc = Document()
    doc.add_heading("AI Resume Screener Report", 0)
    doc.add_paragraph(f"Match Score: {match_score:.2f}%")

    doc.add_heading("Matched Keywords", level=1)
    for kw in sorted(matched):
        doc.add_paragraph(f"✔️ {kw}")

    doc.add_heading("Top 5 Missing Keywords", level=1)
    for kw in sorted(missing)[:5]:
        doc.add_paragraph(f"❌ {kw}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_report(match_score, matched, missing):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="AI Resume Screener Report", ln=True, align="C")

    pdf.set_font("Arial", size=12)
    pdf.ln(10)
    pdf.cell(200, 10, txt=f"Match Score: {match_score:.2f}%", ln=True)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Matched Keywords", ln=True)

    pdf.set_font("Arial", size=12)
    for kw in matched:
        pdf.cell(200, 10, txt=f"- {kw}", ln=True)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Top 5 Missing Keywords", ln=True)

    pdf.set_font("Arial", size=12)
    for kw in list(missing)[:5]:
        pdf.cell(200, 10, txt=f"- {kw}", ln=True)

    pdf_bytes = pdf.output(dest='S').encode('latin1')
    buffer = BytesIO(pdf_bytes)
    return buffer

st.title("🔍 AI Resume Screener")

jd_file = st.file_uploader("Upload Job Description (PDF or DOCX)", type=["pdf", "docx"])
uploaded_resumes = st.file_uploader("📄 Upload One or More Resumes", type=["pdf", "docx"], accept_multiple_files=True)

if jd_file and uploaded_resumes:
    jd_text = extract_text_from_pdf(jd_file) if jd_file.name.endswith(".pdf") else extract_text_from_docx(jd_file)
    if not jd_text.strip():
        st.error("❗ JD file seems empty.")
        st.stop()

    jd_keywords = extract_keywords(jd_text)

    for resume_file in uploaded_resumes:
        resume_text = extract_text_from_pdf(resume_file) if resume_file.name.endswith(".pdf") else extract_text_from_docx(resume_file)
        if not resume_text.strip():
            st.warning(f"⚠️ Skipping {resume_file.name}: File is empty.")
            continue

        resume_keywords = extract_keywords(resume_text)

        categorized_jd_skills = categorize_skills(jd_keywords)
        categorized_resume_skills = categorize_skills(resume_keywords)

        matched_skills = {cat: [] for cat in SKILL_CATEGORIES.keys()}
        missing_skills = {cat: [] for cat in SKILL_CATEGORIES.keys()}

        for category in SKILL_CATEGORIES.keys():
            jd_cat_skills = set(categorized_jd_skills[category])
            resume_cat_skills = set(categorized_resume_skills[category])

            matched_skills[category] = list(jd_cat_skills.intersection(resume_cat_skills))
            missing_skills[category] = list(jd_cat_skills.difference(resume_cat_skills))

        total_jd_skills = sum(len(v) for v in categorized_jd_skills.values())
        total_matched = sum(len(v) for v in matched_skills.values())
        match_score = (total_matched / total_jd_skills * 100) if total_jd_skills else 0

        st.markdown(f"## 📄 Results for: `{resume_file.name}`")
        st.markdown(f"### 🎯 Overall Match Score: {match_score:.2f}%")
        score_meter(match_score)

        for category in SKILL_CATEGORIES.keys():
            st.markdown(f"#### {category} Skills")
            st.markdown(f"- Matched: {matched_skills[category]}")
            st.markdown(f"- Missing (Top 5): {missing_skills[category][:5]}")

        all_matched = [skill for cat_skills in matched_skills.values() for skill in cat_skills]
        all_missing = [skill for cat_skills in missing_skills.values() for skill in cat_skills]

        report = generate_report(match_score, all_matched, all_missing)
        st.download_button("📅 Download DOCX Report", report, file_name=f"{resume_file.name}_report.docx")

        pdf_report = generate_pdf_report(match_score, all_matched, all_missing)
        st.download_button(
            "📅 Download PDF Report",
            data=pdf_report,
            file_name=f"{resume_file.name}_report.pdf",
            mime="application/pdf"
        )
else:
    st.info("👆 Please upload both a Job Description and at least one Resume to proceed.")
